"""Tests for multi-format document parser and integration."""

import io
import docx
import openpyxl
import pytest
from unittest.mock import AsyncMock, patch

from src.services.document_parser import (
    extract_text,
    extract_text_from_csv,
    extract_text_from_docx,
    extract_text_from_xlsx,
)
from src.config import settings


# ─── Parser Unit Tests ────────────────────────────────────────────────────────

def test_extract_text_txt():
    """Test extracting text from a TXT file."""
    txt_bytes = "Hello World!\nThis is standard text.".encode("utf-8")
    result = extract_text(txt_bytes, "doc.txt")
    assert "Hello World!" in result
    assert "This is standard text." in result


def test_extract_text_markdown():
    """Test extracting text from a Markdown file."""
    md_bytes = "# title\n- item 1\n- item 2".encode("utf-8")
    result = extract_text(md_bytes, "doc.md")
    assert "# title" in result
    assert "- item 1" in result


def test_extract_text_csv():
    """Test CSV formatting and skipping empty rows."""
    csv_content = (
        "header1,header2,header3\n"
        "val1, val2 ,val3\n"
        ", ,\n"  # empty row
        "val4,val5,val6\n"
    )
    csv_bytes = csv_content.encode("utf-8-sig")
    result = extract_text_from_csv(csv_bytes)
    
    assert "Row 1: header1 | header2 | header3" in result
    assert "Row 2: val1 | val2 | val3" in result
    assert "Row 3:" not in result  # skipped empty row
    assert "Row 4: val4 | val5 | val6" in result


def test_extract_text_docx():
    """Test DOCX paragraph and table extraction."""
    doc = docx.Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("  Second paragraph with spacing.  ")
    
    # Add a table
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "h1"
    table.cell(0, 1).text = "h2"
    table.cell(0, 2).text = "h3"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    table.cell(1, 2).text = "c"
    
    stream = io.BytesIO()
    doc.save(stream)
    docx_bytes = stream.getvalue()
    
    result = extract_text_from_docx(docx_bytes)
    assert "First paragraph." in result
    assert "Second paragraph with spacing." in result
    assert "Table 1:" in result
    assert "Row 1: h1 | h2 | h3" in result
    assert "Row 2: a | b | c" in result


def test_extract_text_xlsx():
    """Test XLSX sheets, empty sheet/row filtering and custom output format."""
    wb = openpyxl.Workbook()
    
    # Active sheet
    ws = wb.active
    ws.title = "Sheet A"
    ws.cell(row=1, column=1, value="col1")
    ws.cell(row=1, column=2, value="col2")
    ws.cell(row=2, column=1, value="val1")
    ws.cell(row=2, column=2, value="val2")
    ws.cell(row=3, column=1, value=" ") # empty row
    ws.cell(row=3, column=2, value=None)
    ws.cell(row=4, column=1, value="val3")
    
    # Empty sheet
    ws_empty = wb.create_sheet(title="Sheet Empty")
    ws_empty.cell(row=1, column=1, value=None)
    
    stream = io.BytesIO()
    wb.save(stream)
    xlsx_bytes = stream.getvalue()
    
    result = extract_text_from_xlsx(xlsx_bytes)
    assert "Sheet: Sheet A" in result
    assert "Row 1: col1 | col2" in result
    assert "Row 2: val1 | val2" in result
    assert "Row 3:" not in result  # Empty row is skipped
    assert "Row 4: val3" in result
    
    assert "Sheet: Sheet Empty" not in result  # Fully empty sheet is skipped


def test_safety_truncation_limits():
    """Test safety limits are applied and append truncation message."""
    with patch.object(settings, "MAX_EXTRACTED_CHARS", 50):
        text_bytes = ("a" * 100).encode("utf-8")
        result = extract_text(text_bytes, "large.txt")
        assert len(result) <= 50
        assert "[Extraction truncated due to size limit]" in result


def test_safety_file_size_limit():
    """Test files exceeding MAX_FILE_SIZE_MB trigger a ValueError."""
    with patch.object(settings, "MAX_FILE_SIZE_MB", 0.00001): # very small
        text_bytes = ("a" * 200).encode("utf-8")
        with pytest.raises(ValueError) as exc:
            extract_text(text_bytes, "large.txt")
        assert "exceeds safety limit" in str(exc.value)


# ─── API Integration Tests ───────────────────────────────────────────────────

def test_upload_unsupported_format_returns_422(client):
    """Verify that uploading an unsupported format returns HTTP 422."""
    file_payload = {"file": ("image.png", b"fake binary image data", "image/png")}
    response = client.post(
        "/api/v1/documents/upload",
        files=file_payload,
        data={
            "title": "Invalid File",
            "type": "other",
            "process_immediately": "false"
        }
    )
    assert response.status_code == 422
    assert "Unsupported file type" in response.json()["detail"]


def test_replace_unsupported_format_returns_422(client):
    """Verify that replacing file with an unsupported format returns HTTP 422."""
    file_payload = {"file": ("image.png", b"fake binary image data", "image/png")}
    with patch("src.api.v1.documents.get_document_by_id", new_callable=AsyncMock) as mock_get:
        mock_get.return_value = {
            "id": "doc-uuid-1",
            "title": "Doc Title",
            "status": "processed",
            "source_object_key": "doc.pdf"
        }
        response = client.post(
            "/api/v1/documents/doc-uuid-1/replace-file",
            files=file_payload,
            data={"process_immediately": "false"}
        )
        assert response.status_code == 422
        assert "Unsupported file type" in response.json()["detail"]


@pytest.mark.asyncio
async def test_parsing_failure_marks_job_and_doc_failed():
    """Verify that parsing failure sets both document and job status to failed without inserting chunks."""
    from src.services.ingestion_service import process_next_job
    
    claim_mock = {"job_id": "job-uuid-1"}
    
    # Ingestion job points to a corrupted docx file
    job_mock = {
        "job_id": "job-uuid-1",
        "document_id": "doc-uuid-1",
        "source_bucket": "docs",
        "source_object_key": "corrupted.docx",
        "status": "pending",
    }
    
    with patch("src.services.ingestion_service.get_and_claim_next_pending_job", new_callable=AsyncMock) as mock_claim, \
         patch("src.services.ingestion_service.get_job_by_id", new_callable=AsyncMock) as mock_get_job, \
         patch("src.services.ingestion_service.claim_job", new_callable=AsyncMock), \
         patch("src.services.ingestion_service.download_object") as mock_download, \
         patch("src.services.ingestion_service.insert_document_chunks", new_callable=AsyncMock) as mock_insert, \
         patch("src.services.ingestion_service.update_job_status", new_callable=AsyncMock) as mock_update_job, \
         patch("src.services.ingestion_service.update_document_status", new_callable=AsyncMock) as mock_update_doc:
         
        mock_claim.return_value = claim_mock
        mock_get_job.return_value = job_mock
        mock_download.return_value = b"this is corrupt binary that python-docx cannot read"
        
        with pytest.raises(ValueError) as exc:
            await process_next_job()
            
        assert "Failed to parse DOCX document" in str(exc.value)
        
        # Verify job and document statuses are updated to 'failed'
        mock_update_job.assert_called_once_with("job-uuid-1", "failed", exc.value.args[0])
        mock_update_doc.assert_any_call("doc-uuid-1", "failed")
        
        # Verify no document chunks were inserted/created
        mock_insert.assert_not_called()
