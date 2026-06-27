"""Document parsing services for multi-format support."""

import io
import os
import csv
import logging
from typing import Optional
from pypdf import PdfReader

from src.config import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".csv", ".docx", ".xlsx"}
SUPPORTED_FORMATS_ERROR = (
    "Unsupported file type. Supported formats: PDF, TXT, Markdown, CSV, DOCX, XLSX."
)


def _apply_char_limit(text: str) -> str:
    """Helper to safely apply the character extraction limit."""
    if len(text) > settings.MAX_EXTRACTED_CHARS:
        suffix = "\n[Extraction truncated due to size limit]"
        allowed_len = settings.MAX_EXTRACTED_CHARS - len(suffix)
        return text[:allowed_len] + suffix
    return text


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text content from PDF bytes page by page, checking safety limits."""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        text_parts = []
        total_chars = 0
        truncated = False

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                if total_chars + len(page_text) > settings.MAX_EXTRACTED_CHARS:
                    allowed = settings.MAX_EXTRACTED_CHARS - total_chars
                    text_parts.append(page_text[:allowed])
                    truncated = True
                    break
                else:
                    text_parts.append(page_text)
                    total_chars += len(page_text)

        full_text = "\n".join(text_parts)
        if truncated:
            full_text += "\n[Extraction truncated due to size limit]"
        return full_text
    except Exception as e:
        logger.error(f"Error parsing PDF: {e}")
        raise ValueError(f"Failed to parse PDF document: {e}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode raw bytes as UTF-8 plain text, checking safety limits."""
    try:
        text = file_bytes.decode("utf-8")
    except Exception as e:
        logger.error(f"Error decoding text file: {e}")
        raise ValueError(f"Failed to decode text file as UTF-8: {e}")
    return _apply_char_limit(text)


def extract_text_from_markdown(file_bytes: bytes) -> str:
    """Decode raw bytes as UTF-8 markdown text, checking safety limits."""
    try:
        text = file_bytes.decode("utf-8")
    except Exception as e:
        logger.error(f"Error decoding markdown file: {e}")
        raise ValueError(f"Failed to decode markdown file as UTF-8: {e}")
    return _apply_char_limit(text)


def extract_text_from_csv(file_bytes: bytes) -> str:
    """Parse CSV rows into standard line-delimited format, skipping empty rows."""
    try:
        try:
            decoded_str = file_bytes.decode("utf-8-sig")
        except UnicodeDecodeError:
            try:
                decoded_str = file_bytes.decode("utf-8")
            except UnicodeDecodeError as e:
                raise ValueError(f"Failed to decode CSV file as UTF-8: {e}")

        reader = csv.reader(io.StringIO(decoded_str))
        lines = []
        row_count = 0
        truncated = False

        for idx, row in enumerate(reader):
            # Skip fully empty rows
            if not row or not any(cell.strip() for cell in row):
                continue

            row_count += 1
            if row_count > settings.MAX_CSV_ROWS:
                truncated = True
                break

            row_str = " | ".join(cell.strip() for cell in row)
            lines.append(f"Row {idx + 1}: {row_str}")

        full_text = "\n".join(lines)
        if truncated:
            full_text += "\n[Extraction truncated due to size limit]"
        return _apply_char_limit(full_text)
    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Error parsing CSV: {e}")
        raise ValueError(f"Failed to parse CSV document: {e}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract paragraphs and table cells from DOCX file using python-docx."""
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"Error opening DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX document: {e}")

    parts = []

    # 1. Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 2. Extract tables
    for t_idx, table in enumerate(doc.tables, start=1):
        table_lines = [f"Table {t_idx}:"]
        for r_idx, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells]
            row_str = " | ".join(cells)
            table_lines.append(f"Row {r_idx}: {row_str}")
        parts.append("\n".join(table_lines))

    full_text = "\n".join(parts)
    return _apply_char_limit(full_text)


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract sheet rows from XLSX workbook, skipping empty sheets and rows."""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    except Exception as e:
        logger.error(f"Error opening XLSX: {e}")
        raise ValueError(f"Failed to parse XLSX workbook: {e}")

    parts = []

    for sheet in wb.worksheets:
        sheet_lines = []
        row_count = 0
        truncated = False

        for r_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if row_count >= settings.MAX_XLSX_ROWS_PER_SHEET:
                truncated = True
                break

            # Skip fully empty rows
            if not row or not any(
                val is not None and str(val).strip() != "" for val in row
            ):
                continue

            row_count += 1
            cells = [str(val).strip() if val is not None else "" for val in row]
            row_str = " | ".join(cells)
            sheet_lines.append(f"Row {r_idx}: {row_str}")

        if sheet_lines or truncated:
            parts.append(f"Sheet: {sheet.title}")
            parts.extend(sheet_lines)
            if truncated:
                parts.append("[Extraction truncated due to size limit]")

    full_text = "\n".join(parts)
    return _apply_char_limit(full_text)


def extract_text(
    file_bytes: bytes, filename: str, mime_type: Optional[str] = None
) -> str:
    """
    Extract text content from document bytes based on file format.

    Format detection:
    - Prefer file extension.
    - MIME type used as fallback/logging only.
    """
    _, ext = os.path.splitext(filename.lower())
    logger.info(
        f"Extracting text from file '{filename}' with extension '{ext}' and MIME type '{mime_type}'"
    )

    # Check physical file size limit in MB
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > settings.MAX_FILE_SIZE_MB:
        raise ValueError(
            f"File size exceeds safety limit of {settings.MAX_FILE_SIZE_MB} MB."
        )

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    elif ext in (".md", ".markdown"):
        return extract_text_from_markdown(file_bytes)
    elif ext == ".csv":
        return extract_text_from_csv(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    elif ext == ".xlsx":
        return extract_text_from_xlsx(file_bytes)
    else:
        # Fallback to MIME type
        if mime_type == "application/pdf":
            return extract_text_from_pdf(file_bytes)
        elif mime_type == "text/plain":
            return extract_text_from_txt(file_bytes)
        elif mime_type in ("text/markdown", "text/x-markdown"):
            return extract_text_from_markdown(file_bytes)
        elif mime_type in ("text/csv", "application/csv"):
            return extract_text_from_csv(file_bytes)
        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return extract_text_from_docx(file_bytes)
        elif (
            mime_type
            == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ):
            return extract_text_from_xlsx(file_bytes)

        raise ValueError(SUPPORTED_FORMATS_ERROR)
